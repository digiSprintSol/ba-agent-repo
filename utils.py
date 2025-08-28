import io
from io import BytesIO
import pandas as pd
from dataclasses import dataclass
import io
import pandas as pd
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


@dataclass
class Story:
    module: str
    title: str
    description: str
    acceptance_criteria: list

def paginate(items, page_size, page_index):
    start = page_index * page_size
    return items[start:start + page_size]

def make_excel(stories):
    data = []
    for idx, s in enumerate(stories, start=1):
        data.append({
            "#": idx,
            "Module": s.module,
            "Title": s.title,
            "Description": s.description,
            # Add bullet symbol before each AC
            "Acceptance Criteria": "\n".join([f"• {ac}" for ac in s.acceptance_criteria])

        })
    
    df = pd.DataFrame(data)
    buffer = BytesIO()
    with pd.ExcelWriter(buffer, engine="openpyxl") as writer:
        df.to_excel(writer, index=False, sheet_name="UserStories")
    buffer.seek(0)
    return buffer


def make_docx(stories):
    buffer = BytesIO()
    doc = Document()

    # Title
    doc.add_heading("User Stories", level=0)

    for idx, story in enumerate(stories, start=1):
        # Story Title
        doc.add_heading(f"{idx}. {story.title}", level=1)

        # Module
        p = doc.add_paragraph()
        p.add_run("Module: ").bold = True
        p.add_run(story.module)

        # Description
        p = doc.add_paragraph()
        p.add_run("Description: ").bold = True
        p.add_run(story.description)

        # Acceptance Criteria (bullet points)
        doc.add_paragraph("Acceptance Criteria:", style="Heading 2")
        if story.acceptance_criteria:
            for ac in story.acceptance_criteria:
                doc.add_paragraph(ac, style="List Bullet")
        else:
            doc.add_paragraph("N/A")

    doc.save(buffer)
    buffer.seek(0)
    return buffer
