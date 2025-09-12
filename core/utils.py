from dataclasses import dataclass
import pandas as pd
from docx import Document
import re
from io import BytesIO
from openpyxl import Workbook, load_workbook
from openpyxl.styles import Alignment
import os
import streamlit as st
from typing import List, Dict

@dataclass
class Story:
    module: str
    title: str
    description: str
    acceptance_criteria: list

def paginate(items, page_size, page_index): 
    start = page_index * page_size
    return items[start:start + page_size]

def make_excel(stories):
    data = []
    for idx, s in enumerate(stories, start=1):
        data.append({
            "#": idx,
            "Module": s.module,
            "Title": s.title,
            "Description": s.description,
            # Add bullet symbol before each AC
            "Acceptance Criteria": "\n".join([f"• {ac}" for ac in s.acceptance_criteria])

        })
    
    df = pd.DataFrame(data)
    buffer = BytesIO()
    with pd.ExcelWriter(buffer, engine="openpyxl") as writer:
        df.to_excel(writer, index=False, sheet_name="UserStories")
    buffer.seek(0)
    return buffer


def make_docx(stories):
    buffer = BytesIO()
    doc = Document()

    # Title
    doc.add_heading("User Stories", level=0)

    for idx, story in enumerate(stories, start=1):
        # Story Title
        doc.add_heading(f"{idx}. {story.title}", level=1)

        # Module
        p = doc.add_paragraph()
        p.add_run("Module: ").bold = True
        p.add_run(story.module)

        # Description
        p = doc.add_paragraph()
        p.add_run("Description: ").bold = True
        p.add_run(story.description)

        # Acceptance Criteria (bullet points)
        doc.add_paragraph("Acceptance Criteria:", style="Heading 2")
        if story.acceptance_criteria:
            for ac in story.acceptance_criteria:
                doc.add_paragraph(ac, style="List Bullet")
        else:
            doc.add_paragraph("N/A")

    doc.save(buffer)
    buffer.seek(0)
    return buffer


# def make_testcase_excel(test_cases: list, project_name: str, story_module: str) -> BytesIO:
#     """
#     Append test cases to a single Excel file per project.
#     If file does not exist, create it with headers.
#     """
#     df = pd.DataFrame(test_cases)

#     if "testcase_id" in df.columns:
#         df = df.drop_duplicates(subset=["testcase_id"])

#     project_dir = os.path.join("projects", project_name)
#     os.makedirs(project_dir, exist_ok=True)
#     excel_path = os.path.join(project_dir, "testcases.xlsx")

#     headers = [
#        "Testcases ID", "Scenario ID", "Module/Functionality", "TestScenario", "Functional/Integration",
#          "TestCase Title", "Pre-Condition", "Test Data",
#         "Steps to Execute", "Expected Result", "Actual Result", "Status",
#         "Comments", "Priority", "Positive/Negative", "End to End Testing"
#     ]

#     # Load workbook or create new one
#     if os.path.exists(excel_path):
#         wb = load_workbook(excel_path)
#         ws = wb.active
#     else:
#         wb = Workbook()
#         ws = wb.active
#         ws.title = "TestCases"
#         ws.append(headers)

#     scenario_counter = ws.max_row  # continue IDs from last row
#     testcase_counter = ws.max_row

#     for tc in test_cases:
#         scenario_id = tc.get("scenario_id") or f"SC_{scenario_counter:03d}"
#         module = tc.get("module") or story_module or ""
#         test_scenario = tc.get("test_scenario") or ""
#         func_integ = tc.get("functional_integration") or "Functional"
#         testcase_id = tc.get("testcase_id") or tc.get("test_id") or f"TC_{testcase_counter:03d}"
#         testcase_title = tc.get("testcase_title") or tc.get("title") or ""
#         pre_condition = tc.get("pre_condition") or tc.get("preconditions") or ""
#         test_data = tc.get("test_data", "")

#         steps = tc.get("steps_to_execute") or tc.get("steps") or ""
#         if isinstance(steps, list):
#             numbered = []
#             for i, s in enumerate(steps, start=1):
#                 stext = s.strip()
#                 if not stext:
#                     continue
#                 if re.match(r"^\d+\.", stext):  # already numbered
#                     numbered.append(stext)
#                 else:
#                     numbered.append(f"{i}. {stext}")
#             steps = "\n".join(numbered)
#         elif isinstance(steps, str):
#             steps = steps.strip()

#         expected = tc.get("expected_result") or tc.get("expected") or ""
#         actual = tc.get("actual_result", "")
#         status = tc.get("status", "")
#         comments = tc.get("comments", "")
#         priority = tc.get("priority", "Medium")
#         positive_negative = tc.get("positive_negative") or tc.get("positive/negative") or ""
#         end_to_end = tc.get("end_to_end") or "No"

#         row = [
#             testcase_id, scenario_id, module, test_scenario, func_integ,
#             testcase_title, pre_condition, test_data,
#             steps, expected, actual, status, comments, priority,
#             positive_negative, end_to_end
#         ]
#         ws.append(row)

#         # wrap text in steps column
#         steps_cell = ws.cell(row=ws.max_row, column=9)
#         steps_cell.alignment = Alignment(wrapText=True)

#         scenario_counter += 1
#         testcase_counter += 1

#     # widen Steps column
#     ws.column_dimensions['I'].width = 60

#     wb.save(excel_path)

#     # Return buffer for download
#     buf = BytesIO()
#     wb.save(buf)
#     buf.seek(0)
#     return buf

def make_testcase_excel(test_cases: list, project_name: str, story_module: str) -> BytesIO:
    """
    Append test cases to a single Excel file per project.
    If file does not exist, create it with headers.
    Handles sequential ID assignment and avoids duplicates.
    """
    project_dir = os.path.join("projects", project_name)
    os.makedirs(project_dir, exist_ok=True)
    excel_path = os.path.join(project_dir, "testcases.xlsx")

    headers = [
        "Scenario ID", "Module/Functionality", "TestScenario", "Functional/Integration",
        "Testcases ID", "TestCase Title", "Pre-Condition", "Test Data",
        "Steps to Execute", "Expected Result", "Actual Result", "Status",
        "Comments", "Priority", "Positive/Negative", "End to End Testing"
    ]

    # Load workbook or create new one
    if os.path.exists(excel_path):
        wb = load_workbook(excel_path)
        ws = wb.active
    else:
        wb = Workbook()
        ws = wb.active
        ws.title = "TestCases"
        ws.append(headers)

    # ---- find last used IDs ----
    existing_df = pd.DataFrame(ws.values)
    existing_df.columns = existing_df.iloc[0]
    existing_df = existing_df[1:]  # drop header row

    if not existing_df.empty:
        last_tc = existing_df["Testcases ID"].dropna().max()
        last_sc = existing_df["Scenario ID"].dropna().max()
    else:
        last_tc, last_sc = None, None

    def extract_num(val, prefix):
        if isinstance(val, str) and val.startswith(prefix):
            try:
                return int(val.split("_")[1])
            except:
                return 0
        return 0

    max_tc_num = extract_num(last_tc, "TC") if last_tc else 0
    max_sc_num = extract_num(last_sc, "SC") if last_sc else 0

    # ---- process and append ----
    for tc in test_cases:
        # Assign new IDs sequentially
        max_sc_num += 1
        max_tc_num += 1
        scenario_id = f"SC_{max_sc_num:03d}"
        testcase_id = f"TC_{max_tc_num:03d}"

        module = tc.get("module") or story_module or ""
        test_scenario = tc.get("test_scenario") or ""
        func_integ = tc.get("functional_integration") or "Functional"
        testcase_title = tc.get("testcase_title") or tc.get("title") or ""
        pre_condition = tc.get("pre_condition") or tc.get("preconditions") or ""
        test_data = tc.get("test_data", "")

        steps = tc.get("steps_to_execute") or tc.get("steps") or ""
        if isinstance(steps, list):
            numbered = []
            for i, s in enumerate(steps, start=1):
                stext = s.strip()
                if not stext:
                    continue
                if re.match(r"^\d+\.", stext):  # already numbered
                    numbered.append(stext)
                else:
                    numbered.append(f"{i}. {stext}")
            steps = "\n".join(numbered)
        elif isinstance(steps, str):
            steps = steps.strip()

        expected = tc.get("expected_result") or tc.get("expected") or ""
        actual = tc.get("actual_result", "")
        status = tc.get("status", "")
        comments = tc.get("comments", "")
        priority = tc.get("priority", "Medium")
        positive_negative = tc.get("positive_negative") or tc.get("positive/negative") or ""
        end_to_end = tc.get("end_to_end") or "No"

        row = [
            scenario_id, module, test_scenario, func_integ,
            testcase_id, testcase_title, pre_condition, test_data,
            steps, expected, actual, status, comments, priority,
            positive_negative, end_to_end
        ]

        # Check duplicate before appending
        duplicate = False
        for r in ws.iter_rows(min_row=2, values_only=True):
            if (
                r[2] == test_scenario
                and r[5] == testcase_title
                and r[8] == steps
                and r[9] == expected
            ):
                duplicate = True
                break

        if not duplicate:
            ws.append(row)
            # wrap text in steps column
            steps_cell = ws.cell(row=ws.max_row, column=9)
            steps_cell.alignment = Alignment(wrapText=True)

    # widen Steps column
    ws.column_dimensions['I'].width = 60

    wb.save(excel_path)

    # Return buffer for download
    buf = BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf


def render_paginated_table(
    items: List[dict],
    columns: Dict[str, str],
    page_state_key: str,
    page_size: int = 5
):
    """
    Render a paginated table in Streamlit.
    
    Args:
        items (list[dict]): The list of dicts to display.
        columns (dict): Mapping of field -> heading.
        page_state_key (str): Unique key for session state (e.g., "story_page", "tc_page").
        page_size (int): Rows per page.
    """
    if not items:
        st.info("No data available to display.")
        return

    total = len(items)
    total_pages = (total + page_size - 1) // page_size
    current_page = st.session_state.get(page_state_key, 0)

    # Clamp current page
    current_page = max(0, min(current_page, total_pages - 1))
    st.session_state[page_state_key] = current_page

    start = current_page * page_size
    end = start + page_size
    page_items = items[start:end]

    # Build rows for table
    rows = []
    for i, item in enumerate(page_items):
        row = {"#": i + 1 + start}
        for field, heading in columns.items():
            row[heading] = item.get(field, "")
        rows.append(row)

    st.table(rows)

    # Pagination controls
    nav_cols = st.columns(total_pages + 2)
    if nav_cols[0].button("⬅️ Prev", disabled=current_page <= 0, key=f"{page_state_key}_prev"):
        st.session_state[page_state_key] = max(0, current_page - 1)
    for i in range(total_pages):
        if nav_cols[i + 1].button(str(i + 1), key=f"{page_state_key}_{i}"):
            st.session_state[page_state_key] = i
    if nav_cols[-1].button("Next ➡️", disabled=current_page >= total_pages - 1, key=f"{page_state_key}_next"):
        st.session_state[page_state_key] = min(total_pages - 1, current_page + 1)
